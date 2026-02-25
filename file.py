import streamlit as st
import io
import pandas as pd
from pathlib import Path
from docx import Document
import json
import os
import google.generativeai as genai

# ---------------- CONFIGURATION ----------------
API_KEY = os.getenv("GENAI_API_KEY", "AIzaSyAzvmrIDmOZfx3lLqrfXUE1-OXjDEGKg2c")

# ---------------- SETUP GEMINI ----------------
def setup_gemini(api_key: str) -> None:
    genai.configure(api_key=api_key)

# ---------------- UPLOAD IMAGE ----------------
def upload_image_file(file_path: str) -> object:
    path = Path(file_path)
    if not path.exists() or not path.is_file():
        st.error(f"❌ Image file not found: {file_path}")
        return None
    return genai.upload_file(path=path, mime_type="image/png")

# ---------------- EXTRACT RAW CSV ----------------
def extract_table_as_csv(file_obj: object, model_name: str = "gemini-flash-latest") -> str:
    model = genai.GenerativeModel(model_name)
    prompt = """
Extract the table from this image and return it as clean, comma-separated CSV.
• Automatically detect column headers and include them.
• Include all rows and columns as seen.
• Merge multi-line cells into a single line.
• Return only raw CSV text — no markdown, no explanation.
"""
    resp = model.generate_content([prompt, file_obj])
    return resp.text.strip()

# ---------------- PARSE CSV ----------------
def parse_csv(raw_csv: str) -> pd.DataFrame:
    df = pd.read_csv(io.StringIO(raw_csv), sep=',', engine='python', skipinitialspace=True, on_bad_lines='warn')
    return df

# ---------------- SAVE DOCX ----------------
def save_as_word(df: pd.DataFrame, raw_csv: str) -> bytes:
    doc = Document()
    doc.add_heading('Extracted Table Data', 0)
    doc.add_heading('Raw CSV Output', level=1)
    doc.add_paragraph(raw_csv)
    doc.add_paragraph('-------------------------------')
    doc.add_heading('Formatted Table', level=1)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val) if pd.notna(val) else ''

    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()

# ---------------- STREAMLIT UI ----------------
def main():
    st.title("OCR with AI Integration")

    setup_gemini(API_KEY)

    uploaded_file = st.file_uploader("Upload an image file", type=["png", "jpg", "jpeg"])

    if uploaded_file:
        # Save uploaded file temporarily
        st.image(uploaded_file, caption="Uploaded Image Preview", use_container_width=True)

        temp_path = "temp_uploaded_image.png"
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        file_obj = upload_image_file(temp_path)
        os.remove(temp_path)

        if not file_obj:
            st.error("Failed to upload image to Gemini.")
            return

        with st.spinner("Extracting table..."):
            raw_csv = extract_table_as_csv(file_obj)

        if not raw_csv:
            st.error("Failed to extract table.")
            return

        st.subheader("Raw CSV Output")
        st.code(raw_csv)

        try:
            df = parse_csv(raw_csv)
            st.subheader("Parsed Table Preview")
            st.dataframe(df)
        except Exception as e:
            st.error(f"Failed to parse CSV: {e}")
            return

        # CSV download
        csv_bytes = raw_csv.encode('utf-8')
        st.download_button("Download CSV", data=csv_bytes, file_name="extracted_table.csv", mime="text/csv")

        # JSON download
        json_bytes = df.to_json(orient="records", indent=2).encode('utf-8')
        st.download_button("Download JSON", data=json_bytes, file_name="extracted_table.json", mime="application/json")

        # Word download
        word_bytes = save_as_word(df, raw_csv)
        st.download_button("Download Word Document", data=word_bytes, file_name="extracted_table.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    main()
